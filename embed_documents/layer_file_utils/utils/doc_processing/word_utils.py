from docx import Document
from typing import Union
import io

def extract_text_from_word(file_input: Union[str, io.BytesIO]) -> str:
    """
    Extract text from Word document while preserving structure.
    Handles both file paths and file-like objects.
    
    Args:
        file_input (Union[str, io.BytesIO]): Path to the Word file or file-like object
    
    Returns:
        str: Extracted text with preserved structure
    """
    try:
        # Read Word document
        if isinstance(file_input, str):
            doc = Document(file_input)
        else:
            doc = Document(file_input)
        
        # Extract text from paragraphs
        text = ""
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                text += row_text + "\n"
            text += "\n"
        
        return text
                    
    except Exception as e:
        raise Exception(f"Error extracting text from Word: {str(e)}") 